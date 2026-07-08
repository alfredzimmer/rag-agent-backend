"""Structure-aware chunking for the video-transcript-summary .docx compilations.

Each compiled volume concatenates per-video blocks with a rigid template:
H1 video title (often a raw filename); H2 视频来源 (dropped), 大纲, 总结;
H3 sections under 总结 (One-sentence Summary, Takeaways, In-depth Q&A,
Key Words and Tags, Target Audience, Terminology Explanation); then the raw
transcript appended unheadered as 说话人N: HH:MM:SS - HH:MM:SS marker
paragraphs, each followed by that speaker's text. ~80% of the text is
transcript, so packing whole utterances (never splitting inside one unless it
exceeds the cap) is what decides retrieval quality here.

Every chunk gets a context header "{volume} › {video} › {section}" both in the
text (embedded with the content) and as Header 1/2/3 metadata, mirroring the
markdown path in ingest.py. The video label prefers the One-sentence Summary
over the H1 title because H1 is usually a meaningless filename. .docx files
that don't match the template fall back to ingest.py's generic markdown path
via to_markdown().
"""
from __future__ import annotations

import re

from langchain_core.documents import Document

HEADING_STYLE = re.compile(r"^heading([1-9])$")
SPEAKER = re.compile(r"^说话人\s*(\d+)\s*[::]\s*\d[\d:]{3,7}\s*-\s*\d[\d:]{3,7}\s*$")

SUMMARY_TARGET = 1300
QA_TARGET = 1100
TERM_TARGET = 1200
TRANSCRIPT_TARGET = 1400
TRANSCRIPT_CAP = 2000
MIN_ALNUM = 60  # same floor as the markdown chunk filter in ingest.py

DROP_H2 = "视频来源"
OUTLINE_H2 = "大纲"
SUMMARY_H2 = "总结"


def read_paragraphs(path) -> list[tuple[int, str]]:
    """(heading level, text) per non-empty paragraph; level 0 = body text."""
    from docx import Document as DocxDocument  # lazy so the api image needn't carry it

    out = []
    for p in DocxDocument(str(path)).paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name if p.style is not None else "") or ""
        m = HEADING_STYLE.match(style.replace(" ", "").lower())
        out.append((int(m.group(1)) if m else 0, text))
    return out


def is_video_compilation(paras: list[tuple[int, str]]) -> bool:
    return any(
        lvl == 2 and any(k in text for k in (DROP_H2, OUTLINE_H2, SUMMARY_H2))
        for lvl, text in paras
    )


def to_markdown(paras: list[tuple[int, str]]) -> str:
    """Generic fallback: render headings/body as markdown for chunk_markdown."""
    return "\n\n".join(f"{'#' * lvl} {text}" if lvl else text for lvl, text in paras)


def volume_name(stem: str) -> str:
    stem = re.sub(r"_for AI inputs.*$", "", stem, flags=re.I)
    return re.sub(r"\s+", " ", stem).strip(" _-–")


def classify_h3(text: str) -> str:
    t = text.lower()
    if "one-sentence" in t or "one sentence" in t:
        return "one_sentence"
    if "takeaway" in t:
        return "takeaways"
    if "q&a" in t or "q & a" in t or "in-depth" in t:
        return "qa"
    if "key word" in t or "keyword" in t or "tag" in t:
        return "keywords"
    if "audience" in t:
        return "audience"
    if "terminolog" in t:
        return "terminology"
    return "body"


def split_blocks(paras: list[tuple[int, str]]) -> list[tuple[str, list[tuple[int, str]]]]:
    """Split on H1 into per-video blocks; content before the first H1 keeps title ''."""
    blocks: list[tuple[str, list[tuple[int, str]]]] = []
    title, content = "", []
    for lvl, text in paras:
        if lvl == 1:
            if content:
                blocks.append((title, content))
            title, content = text, []
        else:
            content.append((lvl, text))
    if content:
        blocks.append((title, content))
    return blocks


def parse_block(content: list[tuple[int, str]]):
    """Sort a video block's paragraphs into template sections + transcript
    utterances (speaker id, [lines]). The transcript starts at the first
    说话人 marker and runs to the end of the block."""
    sections: dict[str, list[str]] = {
        "outline": [], "one_sentence": [], "takeaways": [], "qa": [],
        "keywords": [], "audience": [], "terminology": [], "body": [],
    }
    transcript: list[tuple[str, list[str]]] = []
    bucket = "body"
    in_transcript = False
    for lvl, text in content:
        m = SPEAKER.match(text)
        if m:
            in_transcript = True
            transcript.append((m.group(1), []))
            continue
        if in_transcript and lvl == 0:
            if transcript:
                transcript[-1][1].append(text)
            continue
        if lvl == 2:
            in_transcript = False
            if DROP_H2 in text:
                bucket = "drop"
            elif OUTLINE_H2 in text:
                bucket = "outline"
            else:
                bucket = "body"
            continue
        if lvl == 3:
            in_transcript = False
            bucket = classify_h3(text)
            continue
        if bucket != "drop":
            sections[bucket].append(text)
    return sections, transcript


def hard_split(text: str, size: int) -> list[str]:
    out = []
    while len(text) > size:
        cut = text.rfind(" ", int(size * 0.6), size)
        if cut == -1:
            cut = size
        out.append(text[:cut].rstrip())
        text = text[cut:].lstrip()
    if text:
        out.append(text)
    return out


def pack(pieces: list[str], target: int, cap: int | None = None) -> list[str]:
    """Greedily join pieces up to ~target chars without splitting a piece,
    unless a single piece exceeds the cap."""
    cap = cap or int(target * 1.5)
    chunks: list[str] = []
    cur = ""
    for piece in pieces:
        piece = piece.strip()
        if not piece:
            continue
        if cur and len(cur) + len(piece) + 1 > target:
            chunks.append(cur)
            cur = ""
        for part in hard_split(piece, cap) if len(piece) > cap else [piece]:
            if cur and len(cur) + len(part) + 1 > target:
                chunks.append(cur)
                cur = part
            else:
                cur = f"{cur}\n{part}" if cur else part
    if cur:
        chunks.append(cur)
    return chunks


def qa_pairs(paras: list[str]) -> list[str]:
    """Group Q&A paragraphs so a question travels with its answer: a paragraph
    ending in ?/？ starts a new pair."""
    pairs: list[list[str]] = []
    for text in paras:
        if text.rstrip().endswith(("?", "？")) or not pairs:
            pairs.append([text])
        else:
            pairs[-1].append(text)
    return ["\n".join(p) for p in pairs]


def chunk_docx(paras: list[tuple[int, str]], *, volume: str) -> list[Document]:
    docs = []
    for title, content in split_blocks(paras):
        sections, transcript = parse_block(content)
        one_sentence = " ".join(sections["one_sentence"]).strip()
        video = one_sentence[:120].rstrip() or title or "untitled"
        summary_parts = (
            sections["one_sentence"] + sections["takeaways"] + sections["outline"]
            + sections["keywords"] + sections["audience"]
        )
        utterances = [
            f"说话人{speaker}: {' '.join(lines)}" for speaker, lines in transcript if lines
        ]
        buckets = [
            ("summary", "Summary", pack(summary_parts, SUMMARY_TARGET)),
            ("qa", "Q&A", pack(qa_pairs(sections["qa"]), QA_TARGET)),
            ("terminology", "Terminology", pack(sections["terminology"], TERM_TARGET)),
            ("body", "Body", pack(sections["body"], SUMMARY_TARGET)),
            ("transcript", "Transcript", pack(utterances, TRANSCRIPT_TARGET, TRANSCRIPT_CAP)),
        ]
        for section_type, label, chunks in buckets:
            for text in chunks:
                if sum(c.isalnum() for c in text) < MIN_ALNUM:
                    continue
                docs.append(
                    Document(
                        page_content=f"{volume} › {video} › {label}\n\n{text}",
                        metadata={
                            "Header 1": volume,
                            "Header 2": video,
                            "Header 3": label,
                            "section_type": section_type,
                            "video_title": title,
                        },
                    )
                )
    return docs
