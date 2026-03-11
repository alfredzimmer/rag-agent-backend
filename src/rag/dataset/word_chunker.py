import re
import json
from docx import Document as DocxDocument
import os
import argparse
import pathlib
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_ollama import ChatOllama

OLLAMA_MODEL = "qwen3:30b-instruct"

def split_word(file: str) -> list[Document]:
    """
    Split the markdown conversion of a DOC file into chunks with size requirements.
    """
    doc = DocxDocument(file)
    # Join all paragraphs into a single text with double newlines to preserve structure
    full_text = "\n\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    # Create a single Document object
    single_doc = Document(page_content=full_text)

    # split raw_splits again with chunk size constraints
    chunk_size = 1024
    chunk_overlap = 256
    sized_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=chunk_overlap, separators=["\n\n", "\n", ".", "?"]
    )

    sized_splits = sized_splitter.split_documents([single_doc])

    return sized_splits


# For debugging purpose 
def format_splits_as_list(splits) -> list[dict]:
    """
    Format the markdown header splits into a well-structured list of dictionaries.

    Args:
        splits: List of Document objects from MarkdownHeaderTextSplitter

    Returns:
        List of dictionaries with headers and content
    """
    formatted_splits = []

    for i, doc in enumerate(splits):
        sanitized_content = doc.page_content.strip()
        split_data = {
            "chunk_id": i,
            "metadata": doc.metadata,
            "content": sanitized_content,
            "char_count": len(doc.page_content)
        }
        formatted_splits.append(split_data)

    return formatted_splits


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Split a DOC file into chunks with size requirements.")
    parser.add_argument("--input", type=str, required=True, help="Path to the DOC folder.")
    parser.add_argument("--output", type=str, default="/srv/shared-data/training-datasets/rag-data/parsed", help="Path to the output JSON folder.")
    args = parser.parse_args()

    INPUT_PATH = args.input
    OUTPUT_PATH = args.output
    for filename in os.listdir(INPUT_PATH):
        if filename.endswith(".docx"):
            FILE_PATH = os.path.join(INPUT_PATH, filename)
            chunks = split_word(FILE_PATH)
            formatted_list = format_splits_as_list(chunks)

            # Optional: Save to JSON file

            filename = os.path.basename(FILE_PATH)          # "{filename}.pdf"
            name_without_ext = os.path.splitext(filename)[0] # "{filename}"
            import json
            output_file = pathlib.Path(f"{OUTPUT_PATH}/{name_without_ext}.json")
            output_file.write_text(json.dumps(formatted_list, indent=2, ensure_ascii=False))
            print(f"\nSaved {len(formatted_list)} chunks to {output_file}")