"""Structure-aware docx chunking against a synthetic template fixture."""
import tempfile
import unittest
from pathlib import Path

try:
    import docx  # noqa: F401
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

if HAVE_DOCX:
    from rag import ingest_docx
    from rag.ingest import process_file

TRANSCRIPT_LINE = (
    "Today we will proceed to Lecture 8, with the theme being Fire Alarm "
    "System Design, one of the most heavily regulated areas in practice. "
)


def build_template_docx(path: Path) -> None:
    d = docx.Document()
    d.add_heading("NA Electrical Design Training L8_1080", level=1)
    d.add_heading("视频来源", level=2)
    d.add_paragraph("user_5809189e/video/2025-9-3/secret-source-url.mp4")
    d.add_heading("大纲", level=2)
    d.add_paragraph("Course Opening and Confirmation (00:00:05)")
    d.add_paragraph("Fire Alarm System Overview (00:01:07)")
    d.add_heading("总结", level=2)
    d.add_heading("One-sentence Summary", level=3)
    d.add_paragraph("This lecture details fire alarm system design and commissioning.")
    d.add_heading("Takeaways", level=3)
    d.add_paragraph("FACP controllers need emergency power supply configuration.")
    d.add_paragraph("Class A loops keep communicating after a single break.")
    d.add_heading("In-depth Q&A", level=3)
    d.add_paragraph("Why must fire alarm systems be interlocked with HVAC systems?")
    d.add_paragraph("HVAC air circulation can accelerate fire spread; duct detectors warn early.")
    d.add_paragraph("What is the core difference between Class A and Class B circuits?")
    d.add_paragraph("Class A uses a loop design; Class B is open-loop with an end-of-line resistor.")
    d.add_heading("Key Words and Tags", level=3)
    d.add_paragraph("Fire Alarm System")
    d.add_heading("Target Audience", level=3)
    d.add_paragraph("Electrical designers working on building projects in North America.")
    d.add_heading("Terminology Explanation", level=3)
    d.add_paragraph("FACP: Fire Alarm Control Panel, the core controller of the system.")
    d.add_paragraph("EOL: End-of-line resistor used to supervise Class B circuits.")
    for i in range(40):
        d.add_paragraph(f"说话人{1 + i % 2}: 00:0{i % 10}:05 - 00:0{i % 10}:59")
        d.add_paragraph(TRANSCRIPT_LINE * 3)
    d.save(str(path))


@unittest.skipUnless(HAVE_DOCX, "python-docx not installed (ingest group)")
class DocxChunkingTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.TemporaryDirectory()
        cls.path = Path(cls.tmp.name) / "Volume 8 – Fire Alarm Systems_for AI inputs_2026-01-12.docx"
        build_template_docx(cls.path)
        cls.docs = process_file(
            cls.path, chunk_size=1800, overlap=250, category="test",
            pdf_parser="auto", fast_pdf_mb=15,
        )

    @classmethod
    def tearDownClass(cls):
        cls.tmp.cleanup()

    def test_template_docx_uses_structured_parser(self):
        self.assertTrue(self.docs)
        self.assertEqual({d.metadata["parser"] for d in self.docs}, {"docx-structured"})

    def test_video_source_section_is_dropped(self):
        for d in self.docs:
            self.assertNotIn("secret-source-url", d.page_content)

    def test_typed_sections_are_produced(self):
        types = {d.metadata["section_type"] for d in self.docs}
        self.assertLessEqual({"summary", "qa", "terminology", "transcript"}, types)

    def test_context_header_prefers_summary_over_filename_title(self):
        for d in self.docs:
            first_line = d.page_content.splitlines()[0]
            self.assertIn("Volume 8 – Fire Alarm Systems › This lecture details", first_line)
            self.assertNotIn("_for AI inputs", first_line)
            self.assertNotIn("L8_1080", first_line)

    def test_qa_answer_travels_with_its_question(self):
        qa = [d for d in self.docs if d.metadata["section_type"] == "qa"]
        chunk = next(d for d in qa if "interlocked with HVAC" in d.page_content)
        self.assertIn("accelerate fire spread", chunk.page_content)

    def test_transcript_chunks_respect_cap_and_keep_speakers(self):
        transcript = [d for d in self.docs if d.metadata["section_type"] == "transcript"]
        self.assertTrue(transcript)
        for d in transcript:
            self.assertLessEqual(len(d.page_content), ingest_docx.TRANSCRIPT_CAP + 200)
            self.assertIn("说话人", d.page_content)
            self.assertNotIn("00:0", d.page_content)  # timestamps stripped

    def test_metadata_keeps_headers_and_stamping(self):
        d = self.docs[0]
        self.assertEqual(d.metadata["Header 1"], "Volume 8 – Fire Alarm Systems")
        self.assertEqual(d.metadata["source_ext"], ".docx")
        self.assertEqual(d.metadata["n_chunks"], len(self.docs))

    def test_plain_docx_falls_back_to_markdown_path(self):
        plain = Path(self.tmp.name) / "Book names.docx"
        d = docx.Document()
        for i in range(30):
            d.add_paragraph(
                f"Reference book {i}: Canadian Electrical Code Handbook, "
                "a practical guide to interpreting installation rules."
            )
        d.save(str(plain))
        docs = process_file(
            plain, chunk_size=1800, overlap=250, category="test",
            pdf_parser="auto", fast_pdf_mb=15,
        )
        self.assertTrue(docs)
        self.assertEqual({x.metadata["parser"] for x in docs}, {"docx-markdown"})


if __name__ == "__main__":
    unittest.main()
